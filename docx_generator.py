"""
docx_generator.py
Fills APSS_Simple_Template.docx directly from JSON dict returned by AI.
No markdown parsing needed - direct placeholder replacement.
"""

import re
from docx import Document


# All template placeholders
ALL_PLACEHOLDERS = [
    "propertyAddress",
    "tenure", "propertyType", "leaseTerm", "councilTaxBand",
    "listedBuildingOrConservationArea", "parkingArrangements",
    "constructionType", "serviceRentGrounds", "sharedOwnershipPercent",
    "materialPropertyNotes",
    "titleStatus", "titleComments",
    "ta6Status", "ta6Comments",
    "ta10Status", "ta10Comments",
    "addInfoStatus", "addInfoComments",
    "epcStatus", "epcComments",
    "searchStatus", "searchComments",
    "leaseInfoStatus", "leaseInfoComments",
    "planningStatus", "planningComments",
    "tenancyStatus", "tenancyComments",
    "probateStatus", "probateComments",
    "hmoStatus", "hmoComments",
    "miNotesStatus", "materialPropertyNotesComments",
    "flagTitleDiscrepancies", "flagRightsOfWay", "flagRestrictiveCovenants",
    "flagPlanningIssues", "flagEpcBelowE", "flagUnregisteredTitle",
    "flagHmoLicensing", "flagLeaseholdInfo", "flagEnvironmentalRisks",
    "flagOccupiersTenancy", "flagSoldAsSeen",
]


def _replace_in_paragraph(para, old: str, new: str) -> bool:
    """Replace placeholder in paragraph, handling split runs."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_text = full.replace(old, new)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""
    return True


def _replace_in_cell(cell, old: str, new: str):
    for para in cell.paragraphs:
        _replace_in_paragraph(para, old, new)
    for table in cell.tables:
        _replace_in_table(table, old, new)


def _replace_in_table(table, old: str, new: str):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, old, new)


def _replace_everywhere(doc: Document, old: str, new: str):
    """Replace placeholder throughout entire document."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, old, new)
    for table in doc.tables:
        _replace_in_table(table, old, new)


def fill_template(template_path: str, field_values: dict, output_path: str) -> bool:
    """
    Fill APSS_Simple_Template.docx with field values from JSON dict.
    field_values keys must match placeholder names (without {{ }}).
    """
    doc = Document(template_path)

    for placeholder in ALL_PLACEHOLDERS:
        template_tag = "{{" + placeholder + "}}"
        value = str(field_values.get(placeholder, "")).strip()
        if not value:
            value = ""
        _replace_everywhere(doc, template_tag, value)

    doc.save(output_path)
    return True


def check_remaining_placeholders(docx_path: str) -> list:
    """Check if any {{ }} placeholders remain unfilled."""
    doc = Document(docx_path)
    remaining = []

    def check(text):
        for m in re.findall(r"\{\{[^}]+\}\}", text):
            remaining.append(m)

    for para in doc.paragraphs:
        check(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                check(cell.text)

    return list(set(remaining))


def json_to_markdown_preview(field_values: dict) -> str:
    """Generate a readable markdown preview from JSON field values."""
    lines = ["# Auction Pack Summary Sheet (APSS)\n"]

    lines.append("## Property Overview\n")
    lines.append(f"**Address:** {field_values.get('propertyAddress', '')}\n")

    overview_fields = [
        ("Tenure", "tenure"),
        ("Property Type", "propertyType"),
        ("Lease Term", "leaseTerm"),
        ("Council Tax Band", "councilTaxBand"),
        ("Listed Building / Conservation Area", "listedBuildingOrConservationArea"),
        ("Parking Arrangements", "parkingArrangements"),
        ("Construction Type", "constructionType"),
        ("Service Charges / Ground Rent", "serviceRentGrounds"),
        ("Shared Ownership %", "sharedOwnershipPercent"),
    ]
    for label, key in overview_fields:
        val = field_values.get(key, "")
        if val:
            lines.append(f"**{label}:** {val}\n")

    mat = field_values.get("materialPropertyNotes", "")
    if mat:
        lines.append(f"\n**Material Information Notes:**\n{mat}\n")

    lines.append("\n## Key Document Status\n")
    doc_rows = [
        ("Title Register & Plan", "titleStatus", "titleComments"),
        ("TA6 Property Info Form", "ta6Status", "ta6Comments"),
        ("TA10 Fixtures & Contents", "ta10Status", "ta10Comments"),
        ("Additional Information Form", "addInfoStatus", "addInfoComments"),
        ("EPC", "epcStatus", "epcComments"),
        ("Searches", "searchStatus", "searchComments"),
        ("Leasehold Information", "leaseInfoStatus", "leaseInfoComments"),
        ("Planning & Certificates", "planningStatus", "planningComments"),
        ("Tenancy Agreements", "tenancyStatus", "tenancyComments"),
        ("Probate or LPA Docs", "probateStatus", "probateComments"),
        ("HMO / Licences / Safety Docs", "hmoStatus", "hmoComments"),
        ("Material Info Notes (Documentation)", "miNotesStatus", "materialPropertyNotesComments"),
    ]
    for label, sk, ck in doc_rows:
        status = field_values.get(sk, "")
        comments = field_values.get(ck, "")
        lines.append(f"**{label}:** {status}")
        if comments:
            lines.append(f"  → {comments[:200]}{'...' if len(comments) > 200 else ''}")
        lines.append("")

    lines.append("\n## Material Flags\n")
    flags = [
        ("Title Discrepancies", "flagTitleDiscrepancies"),
        ("Rights of Way / Access", "flagRightsOfWay"),
        ("Restrictive Covenants", "flagRestrictiveCovenants"),
        ("Known Planning Issues", "flagPlanningIssues"),
        ("EPC < E Rating", "flagEpcBelowE"),
        ("Unregistered Title", "flagUnregisteredTitle"),
        ("HMO / Licencing", "flagHmoLicensing"),
        ("Leasehold Information", "flagLeaseholdInfo"),
        ("Coastal/Flood/Knotweed", "flagEnvironmentalRisks"),
        ("Occupiers / Tenancy", "flagOccupiersTenancy"),
        ("Sold As Seen", "flagSoldAsSeen"),
    ]
    for label, key in flags:
        val = field_values.get(key, "")
        lines.append(f"**{label}:** {val}")

    return "\n".join(lines)